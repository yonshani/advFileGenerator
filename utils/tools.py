import csv
import ctypes
import json
import os
import platform
import logging
from pathlib import Path
from docx import Document


from utils.decorators import log_function_status

os_name = platform.system()
if os_name == 'Windows':
    from send2trash.win.modern import send2trash  # Import send2trash for Windows.

    logging.info("Running on Windows")
elif os_name == 'Darwin':
    logging.info("Running on macOS")
    from send2trash.mac.modern import send2trash  # Import send2trash for macOS.
else:
    # Handle unsupported operating systems.
    logging.info(f"Running on {os_name} - Not supported")

class Tools:

    # Define the default folder path where generated files will be stored.
    default_folder_path = 'generated_files\\'
    # Get the current operating system name.


    def __init__(self):
        pass

    @log_function_status
    def read_json_file(self,file_path) -> None:
        """Reads a JSON file and returns its contents as a dictionary."""
        try:
            with open(file_path, 'r') as json_file:
                # Load the JSON data from the file.
                data = json.load(json_file)
                return data
        except FileNotFoundError:
            # Handle the case where the file is not found.
            logging.debug(f"Error: The file '{file_path}' was not found.")
        except json.JSONDecodeError:
            # Handle invalid JSON file errors.
            logging.debug(f"Error: The file '{file_path}' is not a valid JSON file.")
        except Exception as e:
            # Catch any other exceptions and print an error message.
            logging.debug(f"An error occurred: {e}")

    @log_function_status
    def create_file_with_text(self,folder_path=default_folder_path, file_name=None, text=None, recycle_bin=False) -> None:
        """Creates a text file and writes the specified text into it."""
        full_path = os.path.join(folder_path, file_name)  # Construct the full file path.

        # Check if the folder exists; if not, create it.
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        try:
            # Open the file for writing and write the provided text.
            with open(full_path, 'w') as file:
                file.write(text)
            logging.info(f"File '{full_path}' created successfully.")  # Confirm successful creation.
        except Exception as e:
            # Handle any exceptions that occur during file creation.
            logging.debug(f"An error occurred while creating the file: {e}")

        # If the trash flag is set, move the file to the trash.
        if recycle_bin: send2trash(full_path)

    @log_function_status
    def list_generator(self,items):
        """A generator that yields items from the given list one at a time."""
        for item in items:
            yield item


    @log_function_status
    def reverse_generator(self,lst):
        """A generator that yields items from the given list in reverse order."""
        for item in reversed(lst):
            yield item



    def create_hidden_file_mac(self,folder_path=default_folder_path, file_name=None, text=None, recycle_bin=False):
        """Creates a text file and writes the specified text into it."""
        hiddle_file_name = f".{file_name}"
        full_path = os.path.join(folder_path, hiddle_file_name)  # Construct the full file path.

        # Check if the folder exists; if not, create it.
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        try:
            # Open the file for writing and write the provided text.
            with open(full_path, 'w') as file:
                file.write(text)
            logging.info(f"File '{full_path}' created successfully.")  # Confirm successful creation.
        except Exception as e:
            # Handle any exceptions that occur during file creation.
            logging.debug(f"An error occurred while creating the file: {e}")

        # If the trash flag is set, move the file to the trash.
        if recycle_bin: send2trash(full_path)


    def create_hidden_file_windows(self,folder_path=default_folder_path, file_name=None, text=None, recycle_bin=False):
        # Check if the folder exists; if not, create it.
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        """Creates a hidden file on Windows."""
        full_path = os.path.join(folder_path, file_name)  # Create the file in the current directory
        try:
            # Open the file for writing and write the provided text.
            with open(full_path, 'w') as file:
                file.write(text)
            logging.info(f"File '{full_path}' created successfully.")  # Confirm successful creation.
        except Exception as e:
            # Handle any exceptions that occur during file creation.
            logging.debug(f"An error occurred while creating the file: {e}")

        # Set the file attribute to hidden
        ctypes.windll.kernel32.SetFileAttributesW(full_path, 0x02)  # 0x02 is the hidden attribute
            # If the trash flag is set, move the file to the trash.
        if recycle_bin:
            send2trash(full_path)

        logging.info(f"Hidden file '{full_path}' created successfully.")



    def generate_secret_text(self,suffix,secret_obj):
        if "json" in suffix:
            return secret_obj
        elif "csv" in suffix:
            return [secret_obj,secret_obj,secret_obj]
        elif "pem" in suffix:
            return f"-----BEGIN PRIVATE KEY-----\n{secret_obj['example']}\n-----END PRIVATE KEY-----"
        elif "bat" in suffix:
            return f"@echo off\n{secret_obj['example']}"
        elif "cs" in suffix:
            # Example content for the C# file
            csharp_content = f"""using System;

            namespace HelloWorld
            {{
                class Program
                {{
                    static void Main(string[] args)
                    {{
                        Console.WriteLine("{secret_obj['example']}");
                        Console.WriteLine("Hello, World!");
                    }}
                }}
            }}
            """
            return csharp_content
        return secret_obj["example"]

    def generate_secret_text_formatted(self,suffix,secrets):

        if "pem" in suffix:
            return f"-----BEGIN PRIVATE KEY-----\n{secrets}\n-----END PRIVATE KEY-----"
        elif "bat" in suffix:
            return f"@echo off\n{secrets}"
        elif "cs" in suffix:
            # Example content for the C# file
            csharp_content = f"""using System;

            namespace HelloWorld
            {{
                class Program
                {{
                    static void Main(string[] args)
                    {{
                        Console.WriteLine("{secrets}");
                        Console.WriteLine("Hello, World!");
                    }}
                }}
            }}
            """
            return csharp_content
        return secrets

    def is_json_file(self,file_path):
        """Checks if the file at file_path is a valid JSON file."""
        if not os.path.isfile(file_path):
            logging.debug(f"Error: The path '{file_path}' is not a valid file.")
            return False

        try:
            with open(file_path, 'r') as file:
                json.load(file)  # Try to parse the file as JSON
            return True  # File is a valid JSON
        except json.JSONDecodeError:
            return False  # File is not valid JSON
        except Exception as e:
            logging.debug(f"An error occurred: {e}")
            return False

    def create_json_file_windows(self, folder_path=default_folder_path, file_name=None, data=None, recycle_bin=False, is_hidden=False):
        """Writes a JSON object to a file in a pretty-printed format."""
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        full_path = os.path.join(folder_path, file_name)
        try:
            with open(full_path, 'w') as json_file:
                json.dump(data, json_file, indent=4)  # Pretty-print with an indentation of 4 spaces
            logging.info(f"JSON data written to {full_path} successfully.")
        except Exception as e:
            logging.debug(f"An error occurred while writing to the file: {e}")
        if is_hidden: ctypes.windll.kernel32.SetFileAttributesW(full_path, 0x02)  # 0x02 is the hidden attribute

        if recycle_bin: send2trash(full_path)


    def create_csv_file_windows(self, folder_path=default_folder_path, file_name=None, data=None, recycle_bin=False, is_hidden=False) ->None:
        """Creates a CSV file and writes the given data to it."""
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        full_path = os.path.join(folder_path, file_name)
        try:
            with open(full_path, mode='w', newline='') as csv_file:
                writer = csv.writer(csv_file)

                # Write the header (if needed)
                writer.writerow(data[0].keys())  # Assuming the first dictionary contains the headers

                # Write the data
                for row in data:
                    writer.writerow(row.values())

            logging.info(f"CSV file '{full_path}' created successfully.")
        except Exception as e:
            logging.debug(f"An error occurred while creating the CSV file: {e}")
        if is_hidden: ctypes.windll.kernel32.SetFileAttributesW(full_path, 0x02)  # 0x02 is the hidden attribute
        if recycle_bin: send2trash(full_path)


    def create_hidden_directory(self,path=None) ->None:
        # Define the path for the hidden folder
        formatted_path = Path(path)
        # Create the folder
        formatted_path.mkdir(parents=True, exist_ok=True)

        # Set the folder to hidden
        os.system(f'attrib +h "{formatted_path}"')

    def create_doc_file(self, folder_path=default_folder_path, secret=None,is_hidden=False, recycle_bin=False) ->None:
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        formatted_path = Path(folder_path + secret["name"] + ".docx")  # Update this path as needed
        # Create a new Document
        doc = Document()
        # Add a title
        doc.add_heading('My Document Title', level=1)
        # Add a paragraph
        doc.add_paragraph(secret["example"])
        # Add another paragraph with bold text
        doc.add_paragraph('This is another paragraph with ', 'Normal').add_run('bold text.').bold = True
        # Save the document
        doc.save(formatted_path)

        if recycle_bin: send2trash(formatted_path)
        if is_hidden: os.system(f'attrib +h "{formatted_path}"')